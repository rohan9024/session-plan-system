from flask import Flask, request, render_template, send_file
from openpyxl import load_workbook, Workbook
import os
from io import BytesIO
from docx import Document
from flask import jsonify
import html
from bs4 import BeautifulSoup
import json

app = Flask(__name__)


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file.filename == '':
        return 'No selected file'

    if file:
        document = Document(file)
        content = "<table border='1' id='resultTable'>"

        for table in document.tables:
            for row in table.rows:
                content += "<tr>"
                for cell in row.cells:
                    content += f"<td>{cell.text}</td>"
                content += "</tr>"

        content += "</table>"
        return content
    



def is_bold(run):
    return run.bold



def generate_table(header, rows):
    table_html = "<table border='1'>"
    table_html += "<tr>"

    
    for i, col in enumerate(header):
        if i != 1 and i != 3:
            table_html += f"<th>{col}</th>"
    table_html += "</tr>"

    if rows:
        
        try:
            k = int(header[-1])
            if k > 0:
                cell_contents = [[] for _ in range(len(header) - 2)]  

                for i, row in enumerate(rows):
                    for j, cell in enumerate(row[:1] + row[2:3] + row[4:]): 
                       
                        cell_content = cell.split(',')
                        cell_contents[j].extend(cell_content)

                items_per_cell = max(len(content) // k + (1 if len(content) % k > 0 else 0) for content in cell_contents)

                
                for i in range(k):
                    table_html += "<tr>"
                    for j in range(len(header) - 2):  
                        start_idx = i * items_per_cell
                        end_idx = (i + 1) * items_per_cell
                        table_html += f"<td>{', '.join(cell_contents[j][start_idx:end_idx])}</td>"
                    table_html += "</tr>"

        except ValueError:
            pass  

    table_html += "</table>"
    return table_html








































































@app.route('/process_A', methods=['POST'])
def process_document_A():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file.filename == '':
        return 'No selected file'

    if file:
        document = Document(file)
        content = ""
        current_table_header = []
        current_table_rows = []

        for table in document.tables:
            for row in table.rows:
                first_column_text = row.cells[0].text.strip()
                if first_column_text:
                    if current_table_header and current_table_rows:
                        content += generate_table(current_table_header, current_table_rows)
                        content += "<br>"

                    current_table_header = [cell.text.strip() for cell in row.cells]
                    current_table_rows = []
                else:
                    current_table_rows.append([cell.text.strip() for cell in row.cells])

        if current_table_header and current_table_rows:
            content += generate_table(current_table_header, current_table_rows)

        if content == "":
            return 'No tables found'
        else:
            return content













if __name__ == '__main__':
    app.run(debug=True)
